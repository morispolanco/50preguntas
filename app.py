import streamlit as st
import requests
from docx import Document
from docx.shared import Pt
from time import sleep
import os

# Configuración de la página
st.set_page_config(page_title="Generador de Libros - 24 Preguntas", page_icon="📚")

# Título de la aplicación
st.title("📚 Generador de Libros: 24 Preguntas sobre [Tema]")

# Instrucciones para el usuario
st.write("""
Esta aplicación genera un libro titulado **"24 preguntas sobre [tema]"**. 
Especifica el tema y el público objetivo, y el sistema generará automáticamente 
las preguntas y respuestas estructuradas.
""")

# Entrada del usuario
tema = st.text_input("Tema del libro:", placeholder="Ejemplo: Historia de Roma")
publico = st.selectbox("Público objetivo:", ["General", "Académico", "Principiantes"])

# Botón para iniciar la generación
if st.button("Generar Libro"):
    if not tema:
        st.error("Por favor, ingresa un tema.")
    else:
        # Obtener la API Key desde los Secrets de Streamlit
        api_key = st.secrets["DASHSCOPE_API_KEY"]

        # Crear un documento Word
        doc = Document()
        doc.add_heading(f"24 Preguntas sobre {tema}", level=1)

        # Barra de progreso
        progress_bar = st.progress(0)
        status_text = st.empty()

        for i in range(1, 25):  # Cambiado a 24 preguntas
            # Mostrar progreso
            status_text.text(f"Generando capítulo {i} de 24...")
            progress_bar.progress(i / 24)

            # Construir el prompt para la API
            prompt = f"""
            Genera una pregunta relevante sobre el tema "{tema}" para un público {publico}. 
            Proporciona una respuesta detallada de aproximadamente 500-800 palabras, 
            incluyendo contexto histórico, ejemplos prácticos y datos verificables.
            """

            # Llamar a la API de DashScope
            response = requests.post(
                "https://dashscope-intl.aliyuncs.com/compatible-mode/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "qwen-plus",
                    "messages": [
                        {"role": "system", "content": "Eres un escritor experto en cultura general."},
                        {"role": "user", "content": prompt}
                    ]
                }
            )

            # Procesar la respuesta
            if response.status_code == 200:
                data = response.json()
                contenido = data["choices"][0]["message"]["content"]

                # Agregar el capítulo al documento Word
                doc.add_heading(f"Pregunta {i}: {contenido.split('?')[0]}?", level=2)
                doc.add_paragraph(contenido)
                doc.add_page_break()

                # Simular una pausa entre capítulos
                sleep(1)
            else:
                st.error(f"Error al generar el capítulo {i}. Inténtalo de nuevo.")
                break

        # Agregar el footer al documento Word
        footer = doc.sections[0].footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        run = paragraph.add_run("Copyright Hablemos bien")
        run.font.size = Pt(10)
        run.hyperlink = "https://hablemosbien.org"

        # Guardar el documento Word
        file_name = f"24_preguntas_sobre_{tema.replace(' ', '_')}.docx"
        doc.save(file_name)

        # Permitir descargar el archivo
        with open(file_name, "rb") as file:
            st.download_button(
                label="Descargar Libro",
                data=file,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # Limpiar el archivo generado
        os.remove(file_name)

        # Finalizar la barra de progreso
        status_text.text("¡Libro generado con éxito!")
        progress_bar.progress(1.0)
